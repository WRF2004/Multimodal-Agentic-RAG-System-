"""
Document parsers for various file types.
Handles PDF, images, audio, docx, markdown, etc.
"""

import os
import uuid
import structlog
from typing import Optional

from app.core.interfaces import DocumentParserInterface, Document, OCRInterface, ASRInterface
from app.core.registry import register_component, ComponentRegistry

logger = structlog.get_logger()


@register_component(ComponentRegistry.PARSER, "universal")
class UniversalDocumentParser(DocumentParserInterface):
    """
    Multi-format document parser that dispatches to appropriate handlers.
    Integrates OCR for images and ASR for audio.
    """

    def __init__(
        self,
        ocr: Optional[OCRInterface] = None,
        asr: Optional[ASRInterface] = None,
        ocr_enabled: bool = True,
        asr_enabled: bool = True,
        **kwargs
    ):
        self.ocr = ocr
        self.asr = asr
        self.ocr_enabled = ocr_enabled
        self.asr_enabled = asr_enabled

    def supported_types(self) -> list[str]:
        return ["pdf", "txt", "md", "docx", "png", "jpg", "jpeg", "mp3", "wav", "mp4"]

    async def parse(self, file_path: str, file_type: str = None) -> list[Document]:
        if not os.path.exists(file_path):
            logger.error("file_not_found", path=file_path)
            return []

        if file_type is None:
            file_type = file_path.rsplit(".", 1)[-1].lower()

        try:
            if file_type == "pdf":
                return await self._parse_pdf(file_path)
            elif file_type in ("txt", "md"):
                return await self._parse_text(file_path)
            elif file_type == "docx":
                return await self._parse_docx(file_path)
            elif file_type in ("png", "jpg", "jpeg", "bmp", "tiff"):
                return await self._parse_image(file_path)
            elif file_type in ("mp3", "wav", "flac", "m4a"):
                return await self._parse_audio(file_path)
            else:
                logger.warning("unsupported_file_type", file_type=file_type)
                return []
        except Exception as e:
            logger.error("parse_failed", path=file_path, error=str(e))
            return []

    async def _parse_pdf(self, file_path: str) -> list[Document]:
        from pypdf import PdfReader
        reader = PdfReader(file_path)
        documents = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text and text.strip():
                documents.append(Document(
                    id=str(uuid.uuid4()),
                    content=text.strip(),
                    metadata={
                        "source": file_path,
                        "page": i + 1,
                        "total_pages": len(reader.pages),
                        "file_type": "pdf",
                    },
                    source_type="text",
                ))
            # If page has images and OCR is enabled, try OCR
            elif self.ocr_enabled and self.ocr:
                # For simplicity, log the attempt
                logger.info("pdf_page_no_text_attempting_ocr", page=i + 1)
        return documents

    async def _parse_text(self, file_path: str) -> list[Document]:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()
        if not content.strip():
            return []
        return [Document(
            id=str(uuid.uuid4()),
            content=content,
            metadata={"source": file_path, "file_type": os.path.splitext(file_path)[1][1:]},
            source_type="text",
        )]

    async def _parse_docx(self, file_path: str) -> list[Document]:
        from docx import Document as DocxDocument
        doc = DocxDocument(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        if not paragraphs:
            return []
        return [Document(
            id=str(uuid.uuid4()),
            content="\n".join(paragraphs),
            metadata={"source": file_path, "file_type": "docx"},
            source_type="text",
        )]

    async def _parse_image(self, file_path: str) -> list[Document]:
        if not self.ocr_enabled or not self.ocr:
            logger.warning("ocr_disabled_skipping_image", path=file_path)
            return []
        text = await self.ocr.extract_text(file_path)
        if not text:
            logger.warning("ocr_no_text_extracted", path=file_path)
            return []
        return [Document(
            id=str(uuid.uuid4()),
            content=text,
            metadata={"source": file_path, "file_type": "image", "ocr": True},
            source_type="image",
        )]

    async def _parse_audio(self, file_path: str) -> list[Document]:
        if not self.asr_enabled or not self.asr:
            logger.warning("asr_disabled_skipping_audio", path=file_path)
            return []
        text = await self.asr.transcribe(file_path)
        if not text:
            logger.warning("asr_no_text_transcribed", path=file_path)
            return []
        return [Document(
            id=str(uuid.uuid4()),
            content=text,
            metadata={"source": file_path, "file_type": "audio", "asr": True},
            source_type="audio",
        )]